"""
NLP Utilities for Keyword Extraction
Uses a combination of TF-IDF and RAKE for extracting keywords from project abstracts and documents.
Falls back to simple frequency-based extraction if libraries are not available.
"""

import os
import logging
import re
import json
from typing import List, Tuple, Optional
from collections import Counter
import math
from urllib import request as urllib_request
from urllib import error as urllib_error

# Set up logging
logger = logging.getLogger(__name__)

# Common English stop words
STOP_WORDS = {
    'i', 'me', 'my', 'myself', 'we', 'our', 'ours', 'ourselves', 'you', "you're", "you've", "you'll",
    "you'd", 'your', 'yours', 'yourself', 'yourselves', 'he', 'him', 'his', 'himself', 'she', "she's",
    'her', 'hers', 'herself', 'it', "it's", 'its', 'itself', 'they', 'them', 'their', 'theirs',
    'themselves', 'what', 'which', 'who', 'whom', 'this', 'that', "that'll", 'these', 'those', 'am',
    'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had', 'having', 'do', 'does',
    'did', 'doing', 'a', 'an', 'the', 'and', 'but', 'if', 'or', 'because', 'as', 'until', 'while',
    'of', 'at', 'by', 'for', 'with', 'about', 'against', 'between', 'into', 'through', 'during',
    'before', 'after', 'above', 'below', 'to', 'from', 'up', 'down', 'in', 'out', 'on', 'off', 'over',
    'under', 'again', 'further', 'then', 'once', 'here', 'there', 'when', 'where', 'why', 'how', 'all',
    'each', 'few', 'more', 'most', 'other', 'some', 'such', 'no', 'nor', 'not', 'only', 'own', 'same',
    'so', 'than', 'too', 'very', 's', 't', 'can', 'will', 'just', 'don', "don't", 'should', "should've",
    'now', 'd', 'll', 'm', 'o', 're', 've', 'y', 'ain', 'aren', "aren't", 'couldn', "couldn't", 'didn',
    "didn't", 'doesn', "doesn't", 'hadn', "hadn't", 'hasn', "hasn't", 'haven', "haven't", 'isn', "isn't",
    'ma', 'mightn', "mightn't", 'mustn', "mustn't", 'needn', "needn't", 'shan', "shan't", 'shouldn',
    "shouldn't", 'wasn', "wasn't", 'weren', "weren't", 'won', "won't", 'wouldn', "wouldn't",
    'also', 'however', 'thus', 'therefore', 'hence', 'moreover', 'furthermore', 'although', 'though',
    'even', 'still', 'yet', 'already', 'always', 'never', 'ever', 'often', 'sometimes', 'usually',
    'may', 'might', 'must', 'shall', 'would', 'could', 'need', 'used', 'using', 'use', 'uses',
    'based', 'include', 'includes', 'including', 'included', 'within', 'without', 'well', 'many',
    'much', 'several', 'various', 'different', 'similar', 'like', 'such', 'one', 'two', 'three',
    'first', 'second', 'third', 'new', 'old', 'high', 'low', 'good', 'best', 'better', 'great',
    'important', 'significant', 'main', 'major', 'key', 'primary', 'among', 'across', 'along',
    'around', 'away', 'back', 'behind', 'beside', 'beyond', 'throughout', 'toward', 'towards',
    'upon', 'whether', 'whose', 'show', 'shows', 'shown', 'showing', 'result', 'results',
    'study', 'studies', 'research', 'paper', 'work', 'approach', 'method', 'methods', 'provide',
    'provides', 'proposed', 'propose', 'present', 'presents', 'presented', 'aim', 'aims',
    'objective', 'objectives', 'goal', 'goals', 'focus', 'focuses', 'focused', 'analyze',
    'analysis', 'evaluate', 'evaluation', 'develop', 'development', 'developed', 'implement',
    'implementation', 'implemented', 'achieve', 'achieved', 'demonstrate', 'demonstrated'
}

# Lazy loading of KeyBERT model to avoid slow startup
_kw_model = None
_use_keybert = None

def get_keyword_model():
    """Lazy load KeyBERT model if available"""
    global _kw_model, _use_keybert
    
    if _use_keybert is None:
        try:
            from keybert import KeyBERT
            _kw_model = KeyBERT(model='all-MiniLM-L6-v2')
            _use_keybert = True
            logger.info("KeyBERT model loaded successfully")
        except Exception as e:
            logger.warning(f"KeyBERT not available, using fallback extraction: {e}")
            _use_keybert = False
            _kw_model = None
    
    return _kw_model if _use_keybert else None


def tokenize(text: str) -> List[str]:
    """Simple tokenization"""
    # Convert to lowercase and split on non-alphanumeric
    words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())
    return words


def extract_ngrams(tokens: List[str], n: int = 2) -> List[str]:
    """Extract n-grams from token list"""
    ngrams = []
    for i in range(len(tokens) - n + 1):
        ngram = ' '.join(tokens[i:i+n])
        # Skip if any token is a stop word
        if not any(t in STOP_WORDS for t in tokens[i:i+n]):
            ngrams.append(ngram)
    return ngrams


def extract_keywords_simple(text: str, top_n: int = 10) -> List[str]:
    """
    Simple frequency-based keyword extraction with n-grams.
    Used as fallback when KeyBERT is not available.
    """
    if not text or len(text.strip()) < 50:
        return []
    
    tokens = tokenize(text)
    
    # Filter stop words for unigrams
    unigrams = [t for t in tokens if t not in STOP_WORDS and len(t) > 3]
    
    # Get bigrams
    bigrams = extract_ngrams(tokens, 2)
    
    # Count frequencies
    unigram_counts = Counter(unigrams)
    bigram_counts = Counter(bigrams)
    
    # Normalize similar words (simple stemming - remove common suffixes)
    def normalize(word):
        for suffix in ['ing', 'tion', 'sion', 'ment', 'ness', 'ity', 'ies', 'es', 's', 'ed', 'ly']:
            if word.endswith(suffix) and len(word) > len(suffix) + 3:
                return word[:-len(suffix)]
        return word
    
    # Merge similar unigrams
    normalized_counts = {}
    word_map = {}  # normalized -> best original word
    for word, count in unigram_counts.items():
        norm = normalize(word)
        if norm in normalized_counts:
            normalized_counts[norm] += count
            # Keep the shorter/more common form
            if count > unigram_counts.get(word_map.get(norm, ''), 0):
                word_map[norm] = word
        else:
            normalized_counts[norm] = count
            word_map[norm] = word
    
    # Score bigrams higher (they're more specific)
    all_candidates = []
    
    for norm, count in normalized_counts.items():
        word = word_map[norm]
        all_candidates.append((word, count * 1.0))
    
    for bigram, count in bigram_counts.most_common(top_n * 2):
        all_candidates.append((bigram, count * 2.0))  # Boost bigrams more
    
    # Sort by score and deduplicate
    all_candidates.sort(key=lambda x: x[1], reverse=True)
    
    seen = set()
    keywords = []
    for candidate, score in all_candidates:
        # Skip if any word in candidate was already used
        words_in_candidate = set(candidate.split())
        if not words_in_candidate & seen:
            keywords.append(candidate)
            seen.update(words_in_candidate)
            if len(keywords) >= top_n:
                break
    
    return keywords


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text content from a PDF file"""
    text = ""
    
    # Try PyPDF2 first
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        logger.info(f"PDF has {len(reader.pages)} pages")
        for i, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                    logger.info(f"Page {i+1}: extracted {len(page_text)} chars")
            except Exception as page_err:
                logger.warning(f"Error extracting page {i+1}: {page_err}")
        
        if text.strip():
            logger.info(f"PyPDF2 extracted total {len(text)} chars")
            return text.strip()
    except Exception as e:
        logger.error(f"PyPDF2 error: {e}")
    
    # Try pdfplumber as fallback (better for some PDFs)
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            if text.strip():
                logger.info(f"pdfplumber extracted {len(text)} chars")
                return text.strip()
    except ImportError:
        logger.info("pdfplumber not installed, skipping fallback")
    except Exception as e:
        logger.error(f"pdfplumber error: {e}")
    
    logger.warning("Could not extract text from PDF - may be scanned/image-based")
    return ""


def extract_text_from_docx(file_path: str) -> str:
    """Extract text content from a DOCX file"""
    try:
        from docx import Document
        doc = Document(file_path)
        text_parts = []

        # Main body paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text and paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

        # Table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        text_parts.append(cell_text)

        # Header/footer text
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                if paragraph.text and paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            for paragraph in section.footer.paragraphs:
                if paragraph.text and paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())

        text = "\n".join(text_parts).strip()
        if text:
            return text

        # Fallback: read raw OOXML document.xml
        try:
            import zipfile
            from xml.etree import ElementTree as ET

            with zipfile.ZipFile(file_path, 'r') as archive:
                if 'word/document.xml' in archive.namelist():
                    xml_content = archive.read('word/document.xml')
                    root = ET.fromstring(xml_content)
                    namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    xml_text_nodes = root.findall('.//w:t', namespaces)
                    raw_text = ' '.join(node.text for node in xml_text_nodes if node.text)
                    return raw_text.strip()
        except Exception as xml_err:
            logger.warning(f"DOCX XML fallback failed: {xml_err}")

        # Fallback 2: docx2txt (can handle some edge DOCX structures)
        try:
            import docx2txt
            text_from_docx2txt = (docx2txt.process(file_path) or '').strip()
            if text_from_docx2txt:
                return text_from_docx2txt
        except Exception as docx2txt_err:
            logger.warning(f"docx2txt fallback failed: {docx2txt_err}")

        return ""
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        return ""


def infer_focus_area(title: str, abstract: str) -> str:
    text = f"{title or ''} {abstract or ''}".lower()
    if not text.strip():
        return "General"

    focus_map = [
        ("AI", ["artificial intelligence", "machine learning", "deep learning", "neural", "nlp", "computer vision", "ai"]),
        ("Networking", ["network", "routing", "wireless", "sdn", "tcp", "ip", "lan", "wan", "5g"]),
        ("Security", ["security", "secure", "encryption", "cryptography", "malware", "attack", "vulnerability", "authentication", "authorization"]),
        ("IoT", ["iot", "internet of things", "sensor", "embedded", "smart device"]),
        ("Data Science", ["data science", "analytics", "big data", "data mining", "prediction", "forecast"]),
        ("Software Engineering", ["software", "development", "agile", "testing", "devops", "architecture"]),
        ("Database", ["database", "dbms", "sql", "nosql", "data warehouse"]),
        ("HCI", ["human computer", "usability", "interaction", "ui", "ux"]),
        ("Cloud", ["cloud", "aws", "azure", "gcp", "virtualization", "container", "kubernetes"]),
        ("Mobile", ["mobile", "android", "ios", "smartphone"]),
    ]

    best_area = "General"
    best_score = 0

    for area, keywords in focus_map:
        score = 0
        for keyword in keywords:
            if keyword in text:
                score += 1
        if score > best_score:
            best_score = score
            best_area = area

    return best_area


def extract_text_from_file(file_path: str) -> str:
    """Extract text from a file based on its extension"""
    if not file_path or not os.path.exists(file_path):
        return ""
    
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    elif ext == '.txt':
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error reading text file: {e}")
            return ""
    else:
        logger.warning(f"Unsupported file format: {ext}")
        return ""


def extract_keywords(
    text: str,
    top_n: int = 10,
    keyphrase_ngram_range: Tuple[int, int] = (1, 2),
    stop_words: str = 'english',
    use_mmr: bool = True,
    diversity: float = 0.5
) -> List[str]:
    """
    Extract keywords from text using KeyBERT if available, otherwise fallback to simple extraction.
    
    Args:
        text: The input text to extract keywords from
        top_n: Number of keywords to extract
        keyphrase_ngram_range: Range of n-grams for keywords (min, max)
        stop_words: Language for stop words removal
        use_mmr: Use Maximal Marginal Relevance for diversity
        diversity: Diversity score for MMR (0-1, higher = more diverse)
    
    Returns:
        List of extracted keywords
    """
    if not text or len(text.strip()) < 50:
        return []
    
    model = get_keyword_model()
    
    # Use KeyBERT if available
    if model is not None:
        try:
            keywords_with_scores = model.extract_keywords(
                text,
                keyphrase_ngram_range=keyphrase_ngram_range,
                stop_words=stop_words,
                top_n=top_n,
                use_mmr=use_mmr,
                diversity=diversity
            )
            keywords = [kw[0] for kw in keywords_with_scores]
            logger.info(f"Extracted {len(keywords)} keywords using KeyBERT")
            return keywords
        except Exception as e:
            logger.error(f"KeyBERT extraction failed: {e}")
    
    # Fallback to simple extraction
    logger.info("Using simple keyword extraction (KeyBERT not available)")
    return extract_keywords_simple(text, top_n=top_n)


def extract_keywords_from_project(abstract: str, file_path: Optional[str] = None, top_n: int = 10) -> List[str]:
    """
    Extract keywords from a project's abstract and optionally its uploaded document.
    
    Args:
        abstract: The project abstract
        file_path: Optional path to the uploaded document
        top_n: Number of keywords to extract
    
    Returns:
        List of extracted keywords
    """
    # Combine abstract with document text if available
    combined_text = abstract or ""
    
    if file_path:
        document_text = extract_text_from_file(file_path)
        if document_text:
            # Combine but give more weight to abstract by putting it first
            combined_text = f"{abstract}\n\n{document_text}"
    
    if not combined_text.strip():
        return []
    
    # Extract keywords
    keywords = extract_keywords(combined_text, top_n=top_n)
    
    # Clean and deduplicate keywords
    cleaned_keywords = []
    seen = set()
    for kw in keywords:
        kw_lower = kw.lower().strip()
        if kw_lower and kw_lower not in seen and len(kw_lower) > 2:
            seen.add(kw_lower)
            # Capitalize first letter of each word
            cleaned_keywords.append(kw.title())
    
    return cleaned_keywords[:top_n]


def suggest_keywords(abstract: str, existing_keywords: List[str] = None, top_n: int = 5) -> List[str]:
    """
    Suggest additional keywords based on the abstract that aren't already in existing keywords.
    
    Args:
        abstract: The project abstract
        existing_keywords: List of keywords already assigned
        top_n: Number of suggestions to return
    
    Returns:
        List of suggested keywords
    """
    if not abstract:
        return []
    
    existing_lower = set(kw.lower() for kw in (existing_keywords or []))
    
    # Extract more keywords than needed to filter
    all_keywords = extract_keywords(abstract, top_n=top_n * 2)
    
    # Filter out existing keywords
    suggestions = []
    for kw in all_keywords:
        if kw.lower() not in existing_lower:
            suggestions.append(kw.title())
            if len(suggestions) >= top_n:
                break
    
    return suggestions


def cosine_similarity_from_text(text_a: str, text_b: str) -> float:
    """
    Compute cosine similarity between two text documents using term-frequency vectors.
    Returns score in [0, 1].
    """
    if not text_a or not text_b:
        return 0.0

    tokens_a = [t for t in tokenize(text_a) if t not in STOP_WORDS]
    tokens_b = [t for t in tokenize(text_b) if t not in STOP_WORDS]

    if not tokens_a or not tokens_b:
        return 0.0

    freq_a = Counter(tokens_a)
    freq_b = Counter(tokens_b)

    common_terms = set(freq_a.keys()) & set(freq_b.keys())
    dot_product = sum(freq_a[term] * freq_b[term] for term in common_terms)

    magnitude_a = math.sqrt(sum(value * value for value in freq_a.values()))
    magnitude_b = math.sqrt(sum(value * value for value in freq_b.values()))

    if magnitude_a == 0 or magnitude_b == 0:
        return 0.0

    return dot_product / (magnitude_a * magnitude_b)


def compute_similarity_report(input_text: str, projects, top_k: int = 5) -> dict:
    """
    Compare an input text against a queryset/list of projects and return similarity report.
    """
    if not input_text or len(input_text.strip()) < 20:
        return {
            'similarity_score': 0,
            'top_matches': []
        }

    matches = []
    percent_scores = []

    for project in projects:
        project_text = " ".join([
            project.title or "",
            project.abstract or "",
            project.keywords or ""
        ]).strip()

        if not project_text:
            continue

        similarity = cosine_similarity_from_text(input_text, project_text)
        if similarity <= 0:
            continue

        percent_similarity = round(similarity * 100, 2)
        percent_scores.append(percent_similarity)
        matches.append({
            'project_id': project.id,
            'title': project.title,
            'similarity': percent_similarity
        })

    matches.sort(key=lambda item: item['similarity'], reverse=True)
    top_matches = matches[:top_k]

    if percent_scores:
        top_scores = percent_scores[:top_k]
        avg_top = sum(top_scores) / len(top_scores)
        peak_score = max(percent_scores)
        composite_score = (avg_top * 0.7) + (peak_score * 0.3)
    else:
        avg_top = 0.0
        peak_score = 0.0
        composite_score = 0.0

    return {
        'similarity_score': int(round(composite_score)),
        'top_matches': top_matches,
        'local_details': {
            'average_top_matches': round(avg_top, 2),
            'peak_match': round(peak_score, 2),
            'evaluated_projects': len(percent_scores)
        }
    }


def _extract_first_numeric_score(payload):
    """Find a numeric similarity score using known Winston-style keys only."""
    preferred_keys = {
        'score', 'similarity_score', 'similarity', 'plagiarism_score',
        'overall_score', 'overall_similarity'
    }

    if isinstance(payload, dict):
        for key, value in payload.items():
            if key in preferred_keys and isinstance(value, (int, float)):
                return float(value)
        for value in payload.values():
            found = _extract_first_numeric_score(value)
            if found is not None:
                return found
    elif isinstance(payload, list):
        for item in payload:
            found = _extract_first_numeric_score(item)
            if found is not None:
                return found
    return None


def call_winston_similarity(input_text: str) -> dict:
    """
    Call Winston AI plagiarism API and return a normalized similarity score (0-100).
    Returns {'score': int|None, 'raw': dict|None, 'error': str|None}
    """
    api_key = os.getenv('WINSTON_AI_API_KEY', '').strip()
    api_url = os.getenv('WINSTON_AI_API_URL', 'https://api.gowinston.ai/v2/plagiarism').strip()

    if not api_key:
        return {'score': None, 'raw': None, 'error': 'WINSTON_AI_API_KEY not configured'}

    if not input_text or len(input_text.strip()) < 20:
        return {'score': None, 'raw': None, 'error': 'Insufficient text for Winston check'}

    payload = {
        'text': input_text,
        'content': input_text,
    }

    req = urllib_request.Request(
        url=api_url,
        data=json.dumps(payload).encode('utf-8'),
        headers={
            'Authorization': f'Bearer {api_key}',
            'X-API-Key': api_key,
            'Content-Type': 'application/json',
            'Accept': 'application/json',
            'User-Agent': 'AcademicHub-Backend/1.0'
        },
        method='POST'
    )

    try:
        with urllib_request.urlopen(req, timeout=25) as resp:
            body = resp.read().decode('utf-8')
            parsed = json.loads(body) if body else {}
            try:
                if isinstance(parsed, dict):
                    logger.warning("Winston debug: top-level keys=%s", list(parsed.keys()))
                    logger.warning("Winston debug: raw preview=%s", json.dumps(parsed, ensure_ascii=True)[:800])
                else:
                    logger.warning("Winston debug: non-dict response type=%s", type(parsed))
                    logger.warning("Winston debug: raw preview=%s", json.dumps(parsed, ensure_ascii=True)[:800])
            except Exception:
                logger.warning("Winston debug: failed to serialize response preview")

            score = _extract_first_numeric_score(parsed)
            if score is None:
                logger.warning("Winston response missing similarity score keys. Top-level keys: %s", list(parsed.keys()) if isinstance(parsed, dict) else type(parsed))
                return {'score': None, 'raw': parsed, 'error': 'Could not parse Winston score'}

            # Normalize 0-1 scores to 0-100
            normalized_score = score * 100 if score <= 1 else score
            normalized_score = max(0, min(100, normalized_score))

            return {
                'score': int(round(normalized_score)),
                'raw': parsed,
                'error': None
            }
    except urllib_error.HTTPError as http_err:
        try:
            body = http_err.read().decode('utf-8')
        except Exception:
            body = str(http_err)
        return {'score': None, 'raw': None, 'error': f'Winston HTTP error: {body}'}
    except Exception as err:
        return {'score': None, 'raw': None, 'error': f'Winston request failed: {err}'}


def compute_hybrid_similarity(local_score: int, winston_score: Optional[int], local_top_matches: list) -> dict:
    """
    Combine local and Winston scores into a hybrid result.
    Uses 60% Winston + 40% Local when Winston is available.
    """
    if winston_score is None:
        return {
            'similarity_score': int(local_score),
            'method': 'local_cosine_only',
            'components': {
                'local_score': int(local_score),
                'winston_score': None,
                'weights': {'local': 1.0, 'winston': 0.0}
            },
            'top_matches': local_top_matches
        }

    hybrid = int(round((0.4 * float(local_score)) + (0.6 * float(winston_score))))
    hybrid = max(0, min(100, hybrid))

    return {
        'similarity_score': hybrid,
        'method': 'hybrid_local_winston',
        'components': {
            'local_score': int(local_score),
            'winston_score': int(winston_score),
            'weights': {'local': 0.4, 'winston': 0.6}
        },
        'top_matches': local_top_matches
    }
